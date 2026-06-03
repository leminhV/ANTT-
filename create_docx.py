import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_chapter_4():
    doc = Document()

    # Style functions
    def add_heading(text, level=1):
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            if level == 1:
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0, 0, 0)
            elif level == 2:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 0)
            elif level == 3:
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(0, 0, 0)

    def add_paragraph(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        run.bold = bold
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)

    def add_placeholder(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"\n[ ---> CHỖ TRỐNG: CHÈN ẢNH {text.upper()} TẠI ĐÂY <--- ]\n")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(255, 0, 0) # Màu đỏ để dễ nhận biết
        run.bold = True
        run.italic = True

    # Nội dung Chương 4
    add_heading('CHƯƠNG 4. THIẾT KẾ HỆ THỐNG', level=1)
    add_paragraph('Giai đoạn thiết kế hệ thống là bước chuyển giao quan trọng từ các đặc tả yêu cầu sang các giải pháp kỹ thuật cụ thể. Trong chương này, kiến trúc tổng thể, mô hình cơ sở dữ liệu, cơ chế xử lý đồng thời, và thiết kế Trợ lý ảo AI sẽ được trình bày chi tiết.')

    add_heading('4.1. Kiến trúc công nghệ (Technology Stack) và Mô hình triển khai', level=2)
    add_paragraph('Hệ thống được thiết kế theo mô hình Client-Server hiện đại, sử dụng kiến trúc nguyên khối có tổ chức (Modular Monolith). Công nghệ sử dụng bao gồm:')
    add_bullet('Tầng Frontend (Client): ReactJS kết hợp với Vite, UI Components dựa trên Tailwind CSS.')
    add_bullet('Tầng Backend (Server): Xây dựng trên nền tảng NestJS (Node.js framework) hỗ trợ OOP và TypeScript.')
    add_bullet('Tầng Cơ sở dữ liệu: MySQL 8.0 (InnoDB) hỗ trợ đầy đủ Transaction ACID và Row-level Locking, giao tiếp qua Prisma ORM.')
    
    add_paragraph('\nMô hình triển khai vật lý (Deployment Topology):')
    add_bullet('Hệ thống được triển khai trên máy chủ Linux ảo hóa (VPS).')
    add_bullet('Nginx đóng vai trò Reverse Proxy (Port 80/443), điều hướng traffic vào Frontend tĩnh và Backend API.')
    add_bullet('Backend NestJS được đóng gói thông qua Docker (Dockerfile) để đảm bảo tính nhất quán giữa các môi trường.')
    add_placeholder('Sơ đồ triển khai hệ thống (Deployment Diagram)')

    add_heading('4.2. Thiết kế cơ sở dữ liệu (ERD)', level=2)
    add_placeholder('Sơ đồ thực thể kết nối (ERD)')
    add_heading('4.2.1. Các thực thể chính và thuộc tính chi tiết', level=3)
    add_bullet('Users: Quản lý tài khoản (user_id, email, password_hash, role, status).')
    add_bullet('Rooms: Quản lý phòng Lab (room_id, name, capacity, status).')
    add_bullet('Equipments: Quản lý thiết bị (equipment_id, device_name, room_id, status, row_version).')
    add_bullet('Bookings: Quản lý phiên đặt lịch (booking_id, user_id, room_id, start_time, end_time, status).')
    add_bullet('SystemSettings: Lưu trữ cấu hình động Key-Value (thời gian ca tối thiểu, buffer_time).')

    add_heading('4.2.2. Schema SQL: bảng, khóa, index, ràng buộc toàn vẹn', level=3)
    add_paragraph('Các bảng được liên kết chặt chẽ thông qua các Ràng buộc khóa ngoại. Hệ thống tối ưu hóa truy vấn bằng B-Tree Index trên các cột thường xuyên được truy xuất như Bookings.start_time. Khóa ngoại sử dụng ON DELETE RESTRICT để tránh tạo ra dữ liệu mồ côi (Orphan Data).')

    add_heading('4.3. Phân tích và giải quyết xung đột lịch (Concurrency)', level=2)
    add_heading('4.3.1. Kịch bản Race Condition (Lost Update Anomaly)', level=3)
    add_paragraph('Nếu hệ thống chỉ sử dụng thao tác CRUD thông thường, hiện tượng Race Condition sẽ xảy ra khi hai sinh viên cùng gửi yêu cầu đặt một phòng Lab trong cùng một phần nghìn giây. Cả hai luồng xử lý đều đọc được trạng thái "Phòng trống" và cùng ghi dữ liệu, dẫn đến lỗi Double Booking (Đụng độ lịch).')
    add_placeholder('Sơ đồ tuần tự (Sequence Diagram) mô tả kịch bản Race Condition')

    add_heading('4.3.2. Giải pháp SELECT...FOR UPDATE trong MySQL InnoDB', level=3)
    add_paragraph('Nhóm áp dụng cơ chế Khóa bi quan (Pessimistic Locking) qua Prisma Interactive Transaction:')
    add_bullet('Gom câu lệnh kiểm tra (SELECT) và tạo mới (INSERT) vào một Giao dịch (Transaction).')
    add_bullet('Sử dụng lệnh SELECT ... FOR UPDATE để khóa dòng dữ liệu vật lý.')
    add_bullet('Luồng gửi sau sẽ bị Database bắt chờ cho đến khi luồng trước commit transaction. Nếu phát hiện lịch đã bị chiếm, luồng sau sẽ bị từ chối an toàn (Conflict 409).')
    add_placeholder('Ảnh chụp màn hình Test bằng JMeter (Load Testing chứng minh 0 lỗi Double Booking)')

    add_heading('4.4. Thiết kế phân quyền RBAC và ma trận phân quyền', level=2)
    add_paragraph('Hệ thống áp dụng mô hình phân quyền dựa trên vai trò (Role-Based Access Control) với 4 vai trò: Student, Instructor, Technician, Admin. Các giới hạn quyền hạn được ánh xạ theo nguyên tắc Least Privilege (Quyền hạn tối thiểu).')
    add_placeholder('Bảng Ma trận phân quyền (Role Matrix)')

    add_heading('4.5. Thiết kế luồng xác thực JWT (Access + Refresh Token)', level=2)
    add_paragraph('Quy trình xác thực sử dụng bộ đôi JSON Web Token:')
    add_bullet('Access Token (15 phút): Được lưu trong memory của Frontend, gửi qua Authorization header trong mọi request.')
    add_bullet('Refresh Token (7 ngày): Được lưu an toàn dưới dạng HttpOnly Cookie để chống XSS, tự động dùng để lấy Access Token mới khi hết hạn.')
    add_placeholder('Sơ đồ luồng hoạt động của Refresh Token Rotation')

    add_heading('4.6. Triển khai bảo mật trong hệ thống', level=2)
    add_paragraph('Toàn bộ mật khẩu được băm một chiều (Hash) bằng thuật toán bcrypt. Api được bảo vệ bằng NestJS Global Validation Pipe để chống SQL Injection. Các HTTP Security Headers (qua Helmet) như X-Frame-Options, X-XSS-Protection được bật đầy đủ. Lỗi IDOR được khắc phục triệt để bằng cách luôn đối chiếu user_id trong JWT với chủ sở hữu bản ghi.')

    add_heading('4.7. Thiết kế giao diện người dùng (UI/UX Design)', level=2)
    add_paragraph('Giao diện tuân thủ triết lý Flat Design, sử dụng hệ thống lưới 12-column, Typography hiện đại (Inter) và tông màu Trust Blue làm chủ đạo.')
    add_placeholder('Ảnh chụp màn hình Dashboard và Giao diện Đặt lịch (Calendar View)')
    add_placeholder('Sơ đồ luồng màn hình (Screen Flow / Site Map)')

    add_heading('4.8. Thiết kế Module Trợ lý ảo AI Client-side (Zero Cost NLP)', level=2)
    add_paragraph('Đây là phân hệ vệ tinh đột phá của hệ thống, cung cấp tính năng chat tự nhiên để thực hiện thao tác (Ví dụ: "Hãy duyệt tất cả đơn phòng A").')
    add_bullet('Cơ chế hoạt động: Phân tích cú pháp dựa trên Keyword và Regular Expression (Regex) chạy 100% tại Client-side.')
    add_bullet('Ưu điểm: Không phụ thuộc vào API bên thứ ba (như OpenAI API), đảm bảo thời gian phản hồi tức thì (<10ms), bảo mật hoàn toàn dữ liệu chat của người dùng và tối ưu hóa 100% chi phí vận hành (Zero Cost).')
    add_placeholder('Sơ đồ luồng thuật toán của Trợ lý AI (Input -> Intent Matcher -> Action Dispatcher)')
    add_placeholder('Ảnh chụp màn hình đang chat với Trợ lý ảo AI trên giao diện')

    # Lưu file
    file_path = os.path.join(os.getcwd(), 'Chuong4_ThietKeHeThong.docx')
    doc.save(file_path)
    print(f"Đã tạo thành công file: {file_path}")

if __name__ == '__main__':
    create_chapter_4()
